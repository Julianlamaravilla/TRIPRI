#!/usr/bin/env python3
"""Renderiza el cuestionario markdown a .docx conservando negritas, tablas,
citas, listas y los marcadores [Answer]: como campos rellenables."""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC, DST = sys.argv[1], sys.argv[2]

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(6)

def shade(p, color):
    el = OxmlElement('w:shd'); el.set(qn('w:val'), 'clear'); el.set(qn('w:fill'), color)
    p._p.get_or_add_pPr().append(el)

def inline(par, text, bold=False):
    """Procesa **negrita**, *cursiva*, `código` y ~~tachado~~."""
    for tok in re.split(r'(\*\*.+?\*\*|`[^`]+`|(?<!\*)\*(?!\*)[^*]+?\*(?!\*))', text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = par.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
            r = par.add_run(tok[1:-1].replace('`', '')); r.italic = True
        else:
            r = par.add_run(tok)
        if bold:
            r.bold = True
    return par

def answer_field(indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run('[Answer]: '); r.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x6F, 0xB0)
    r2 = p.add_run('_' * 70); r2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    shade(p, 'F2F7FC')

lines = open(SRC, encoding='utf-8').read().split('\n')
i = 0
while i < len(lines):
    ln = lines[i]

    # ---- tablas ----
    if ln.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s:|-]+\|$', lines[i + 1]):
        rows = []
        while i < len(lines) and lines[i].startswith('|'):
            if not re.match(r'^\|[\s:|-]+\|$', lines[i]):
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
            i += 1
        cols = max(len(r) for r in rows)
        t = doc.add_table(rows=0, cols=cols)
        t.style = 'Light Grid Accent 1'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for ci in range(cols):
                txt = row[ci] if ci < len(row) else ''
                par = cells[ci].paragraphs[0]
                par.paragraph_format.space_after = Pt(2)
                inline(par, txt, bold=(ri == 0))
                for r in par.runs:
                    r.font.size = Pt(9)
        doc.add_paragraph()
        continue

    # ---- separador ----
    if ln.strip() == '---':
        p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
        bd = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6'); b.set(qn('w:color'), 'CCCCCC')
        bd.append(b); pPr.append(bd)
        i += 1; continue

    # ---- encabezados ----
    m = re.match(r'^(#{1,4}) (.+)$', ln)
    if m:
        lvl, txt = len(m.group(1)), m.group(2)
        h = doc.add_heading(level=min(lvl, 4))
        h.paragraph_format.space_before = Pt(14 if lvl <= 2 else 10)
        inline(h, txt)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0x14, 0x3D, 0x66)
        i += 1; continue

    # ---- cita / bloque destacado ----
    if ln.startswith('>'):
        block = []
        while i < len(lines) and (lines[i].startswith('>') or (block and lines[i].strip() == '')):
            if lines[i].strip() == '':
                if i + 1 < len(lines) and not lines[i + 1].startswith('>'):
                    break
                block.append('')
            else:
                block.append(re.sub(r'^>\s?', '', lines[i]))
            i += 1
        for b in block:
            if not b.strip():
                continue
            hm = re.match(r'^(#{1,4}) (.+)$', b)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            inline(p, hm.group(2) if hm else re.sub(r'^[-*] ', '• ', b), bold=bool(hm))
            if hm:
                for r in p.runs:
                    r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x8A, 0x2A, 0x1A)
            shade(p, 'FFF6E5')
        continue

    # ---- campo [Answer]: ----
    if '`[Answer]:`' in ln:
        pre = ln.split('`[Answer]:`')[0].strip()
        if pre:
            inline(doc.add_paragraph(), re.sub(r'^[-*] ', '• ', pre))
        answer_field(0.25 if ln.strip().startswith(('-', '*')) else 0.0)
        i += 1; continue

    # ---- listas ----
    m = re.match(r'^(\s*)[-*] (.+)$', ln)
    if m:
        depth = len(m.group(1)) // 2
        p = doc.add_paragraph(style='List Bullet' if depth == 0 else 'List Bullet 2')
        p.paragraph_format.space_after = Pt(3)
        inline(p, m.group(2))
        i += 1; continue

    m = re.match(r'^(\s*)(\d+)\. (.+)$', ln)
    if m:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(3)
        inline(p, m.group(3))
        i += 1; continue

    # ---- párrafo ----
    if ln.strip():
        inline(doc.add_paragraph(), ln.strip())
    i += 1

doc.save(DST)
print(f'OK → {DST}')
